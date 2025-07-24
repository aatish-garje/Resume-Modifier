import streamlit as st
import docx
import pdfplumber
import io
import re
import spacy
from openai import OpenAI
from fuzzywuzzy import fuzz
import spacy.cli
import os
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import spacy.cli
    spacy.cli.download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# --- Setup ---
st.set_page_config(page_title="AI Resume Optimizer", layout="centered")
st.title("🧠 AI Resume Optimizer (Skill Match & Resume Enhancer)")

api_key = st.text_input("🔑 OpenRouter API Key", type="password")
client = OpenAI(api_key=api_key, base_url="https://openrouter.ai/api/v1") if api_key else None

# --- Load NLP ---
@st.cache_resource
def load_nlp():
    return spacy.load("en_core_web_sm")
nlp = load_nlp()

# --- Extract Keywords ---
def extract_keywords(text):
    doc = nlp(text)
    phrases = set()
    for chunk in doc.noun_chunks:
        phrase = chunk.text.strip().lower()
        if 2 <= len(phrase.split()) <= 6:
            phrases.add(phrase)
    return sorted(phrases)

# --- Text Extract ---
def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

def get_text(file):
    if file.name.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif file.name.endswith(".docx"):
        return extract_text_from_docx(file)
    return ""

def write_to_docx(text):
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line.strip())
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# --- Fuzzy Skill Matching ---
def fuzzy_skill_match(jd_skills, resume_skills, threshold=80):
    matched = []
    for js in jd_skills:
        for rs in resume_skills:
            if fuzz.partial_ratio(js, rs) >= threshold:
                matched.append(js)
                break
    return list(set(matched))

def calculate_match_score(jd_skills, resume_skills):
    matched = fuzzy_skill_match(jd_skills, resume_skills)
    return round(100 * len(matched) / len(jd_skills), 1) if jd_skills else 0, matched

# --- GPT Bullet Generator ---
def generate_bullets(client, keywords, jd_text):
    prompt = f"""
You are a professional resume writer. Based on the job description and skills below, write 7–8 impactful resume bullet points that highlight strong technical contributions. Use real achievements, quantifiable impact, and make it ATS-friendly.

Skills: {', '.join(keywords)}
Job Description: {jd_text[:1200]}
Return only bullet points.
"""
    try:
        response = client.chat.completions.create(
            model="openai/gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=700,
            temperature=0.4
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"⚠️ API Error: {e}"

# --- Resume Modifiers ---
def inject_into_resume(resume_text, bullets, section="Experience"):
    pattern = re.compile(section, re.IGNORECASE)
    match = pattern.search(resume_text)
    if match:
        idx = match.end()
        return resume_text[:idx] + "\n" + bullets + "\n" + resume_text[idx:]
    return bullets + "\n\n" + resume_text

def inject_skills_section(resume_text, skills):
    skills_text = "\nSkills:\n" + "\n".join(f"• {kw.title()}" for kw in skills)
    if "skills" in resume_text.lower():
        return resume_text + "\n" + skills_text
    else:
        return resume_text + "\n\n" + skills_text

# --- Upload Resume ---
st.subheader("📤 Upload Resume")
resume_file = st.file_uploader("Choose resume (.pdf or .docx)", type=["pdf", "docx"])

# --- Job Description ---
st.subheader("📝 Job Description")
jd_method = st.radio("Provide JD as:", ["Upload PDF", "Upload DOCX", "Paste Text"])
jd_text = ""

if jd_method == "Upload PDF":
    jd_file = st.file_uploader("Upload JD PDF", type=["pdf"], key="jd_pdf")
    if jd_file:
        jd_text = extract_text_from_pdf(jd_file)
elif jd_method == "Upload DOCX":
    jd_file = st.file_uploader("Upload JD DOCX", type=["docx"], key="jd_docx")
    if jd_file:
        jd_text = extract_text_from_docx(jd_file)
elif jd_method == "Paste Text":
    jd_text = st.text_area("Paste the job description")

# --- Skill Detection ---
st.subheader("🔍 Skills / Keywords")
auto_detect = st.checkbox("✨ Auto-detect from JD", value=True)
manual_keywords = st.text_input("✍️ Add skills manually (comma-separated)", placeholder="Python, SAP, NLP, Data Analysis")

jd_skills = extract_keywords(jd_text) if auto_detect and jd_text else []
manual_list = [kw.strip().lower() for kw in manual_keywords.split(",") if kw.strip()]
jd_skills = list(set(jd_skills + manual_list))

if jd_skills:
    st.markdown("🧠 **Extracted JD Skills:**")
    st.markdown("✅ " + " &nbsp;•&nbsp; ".join([kw.title() for kw in jd_skills]))

# --- Resume Optimization ---
if st.button("🚀 Optimize Resume"):
    if not resume_file or not jd_text.strip():
        st.warning("Please upload both resume and JD.")
    elif not api_key:
        st.error("⚠️ OpenRouter API key is missing.")
    else:
        resume_text = get_text(resume_file)
        if not resume_text.strip():
            st.error("❌ Could not read resume.")
        else:
            resume_skills = extract_keywords(resume_text)
            score, matched_skills = calculate_match_score(jd_skills, resume_skills)
            unmatched = list(set(jd_skills) - set(matched_skills))

            st.subheader("📊 Skill Match Score")
            st.success(f"✅ Skill Match Score: {score}% match between JD and your resume")

            st.markdown("🔍 **Matched Skills:**")
            st.markdown("✔️ " + " &nbsp;•&nbsp; ".join([kw.title() for kw in matched_skills]))

            if unmatched:
                st.markdown("❌ **Missing Skills from Resume:**")
                st.markdown("⚠️ " + " &nbsp;•&nbsp; ".join([kw.title() for kw in unmatched]))

            st.info("⏳ Generating enhanced resume bullets...")
            bullets = generate_bullets(client, unmatched, jd_text)

            st.subheader("✅ GPT-Generated Bullet Points")
            st.code(bullets)

            final_resume = inject_into_resume(resume_text, bullets, section="Experience")
            final_resume = inject_skills_section(final_resume, jd_skills)

            docx_data = write_to_docx(final_resume)

            st.download_button(
                label="📥 Download Optimized Resume (.docx)",
                data=docx_data,
                file_name="optimized_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
